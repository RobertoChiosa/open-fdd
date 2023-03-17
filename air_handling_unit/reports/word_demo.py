# Author:       Roberto Chiosa
# Copyright:    Roberto Chiosa, © 2023
# Email:        roberto.chiosa@polito.it
#
# Created:      16/03/23
# Script Name:  word_demo.py
# Path:         air_handling_unit/reports
#
# Script Description:
#
#
# Notes:
from docx import Document
from docx.shared import Inches

if __name__ == '__main__':

    document = Document()

    document.add_heading('Document Title', 0)

    p = document.add_paragraph('A plain paragraph having some ')
    p.add_run('bold').bold = True
    p.add_run(' and some ')
    p.add_run('italic.').italic = True

    document.add_heading('Heading, level 1', level=1)
    document.add_paragraph('Intense quote', style='Intense Quote')

    document.add_paragraph(
        'first item in unordered list', style='List Bullet'
    )
    document.add_paragraph(
        'first item in ordered list', style='List Number'
    )

    document.add_picture('../images/ahu_operating_modes.png', width=Inches(1.25))

    records = (
        (3, '101', 'Spam'),
        (7, '422', 'Eggs'),
        (4, '631', 'Spam, spam, eggs, and spam')
    )

    table = document.add_table(rows=1, cols=3)
    hdr_cells = table.rows[0].cells
    hdr_cells[0].text = 'Qty'
    hdr_cells[1].text = 'Id'
    hdr_cells[2].text = 'Desc'
    for qty, idx, desc in records:
        row_cells = table.add_row().cells
        row_cells[0].text = str(qty)
        row_cells[1].text = idx
        row_cells[2].text = desc

    document.add_page_break()

    document.save('../final_report/word_demo.docx')
